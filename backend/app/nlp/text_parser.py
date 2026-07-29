import io
import re
import logging
import httpx
from typing import Optional

logger = logging.getLogger("CampusOS.NLP.TextParser")

def parse_pdf_bytes(content: bytes) -> str:
    """Extracts raw text from PDF file bytes using pypdf."""
    try:
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(content))
        extracted_pages = []
        for page in reader.pages:
            t = page.extract_text()
            if t:
                extracted_pages.append(t)
        return "\n".join(extracted_pages).strip()
    except Exception as e:
        logger.warning(f"pypdf extraction failed: {e}")
        return content.decode("utf-8", errors="ignore").strip()

def parse_docx_bytes(content: bytes) -> str:
    """Extracts raw text from DOCX file bytes using python-docx."""
    try:
        import docx
        doc = docx.Document(io.BytesIO(content))
        full_text = []
        for para in doc.paragraphs:
            if para.text:
                full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        full_text.append(cell.text)
        return "\n".join(full_text).strip()
    except Exception as e:
        logger.warning(f"python-docx extraction failed: {e}")
        return content.decode("utf-8", errors="ignore").strip()

def clean_scraped_html(html_content: str) -> str:
    """Strips HTML tags, scripts, and styles to leave readable plain text."""
    # Remove script and style tags
    cleaned = re.sub(r'<(script|style).*?>.*?</\1>', ' ', html_content, flags=re.DOTALL | re.IGNORECASE)
    # Remove HTML tags
    cleaned = re.sub(r'<.*?>', ' ', cleaned)
    # Collapse whitespace
    cleaned = re.sub(r'\s+', ' ', cleaned)
    return cleaned.strip()

async def fetch_url_text(url: str) -> str:
    """Fetches web page content from URL and returns clean plain text."""
    try:
        async with httpx.AsyncClient(timeout=10.0, follow_redirects=True) as client:
            headers = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) CampusOS-AI/1.0"}
            resp = await client.get(url, headers=headers)
            if resp.status_code == 200:
                return clean_scraped_html(resp.text)
    except Exception as e:
        logger.warning(f"URL scraping failed for {url}: {e}")
    return ""

def parse_document_input(
    filename: Optional[str] = None,
    file_bytes: Optional[bytes] = None,
    raw_text: Optional[str] = None
) -> str:
    """Central document parser resolving bytes or text into plain text."""
    if file_bytes and filename:
        fn = filename.lower()
        if fn.endswith(".pdf"):
            text = parse_pdf_bytes(file_bytes)
            if text:
                return text
        elif fn.endswith(".docx") or fn.endswith(".doc"):
            text = parse_docx_bytes(file_bytes)
            if text:
                return text
        else:
            return file_bytes.decode("utf-8", errors="ignore").strip()
    
    if raw_text:
        return raw_text.strip()
    
    return ""
